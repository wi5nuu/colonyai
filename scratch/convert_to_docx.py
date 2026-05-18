import os
import re
import sys

# Auto-install python-docx if not installed
try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    import subprocess
    print("Installing required python-docx package for high-fidelity conversion...")
    try:
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=True)
        import docx
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement, parse_xml
        from docx.oxml.ns import nsdecls, qn
    except Exception as e:
        print(f"Error installing python-docx: {e}")
        print("Please run: pip install python-docx")
        sys.exit(1)

# Helper function to set cell background shading
def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

# Helper function to set cell margins (padding)
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

# Helper function to add borders to a table
def set_table_borders(table, color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

# Inline markdown formatter for Bold and Italic text
def add_formatted_text(paragraph, text):
    # Regex to find bold **text** or italic _text_
    parts = re.split(r'(\*\*.*?\*\*|_.*?_)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('_') and part.endswith('_'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found!")
        return
        
    print(f"Reading {md_path}...")
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    doc = docx.Document()
    
    # Configure page margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Configure Normal Style (Calibri 11pt, 1.15 line spacing)
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Calibri'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    
    in_code_block = False
    code_lines = []
    
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # --- Handle Code / Wireframe Blocks ---
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block, render it in a boxed callout table
                in_code_block = False
                box_table = doc.add_table(rows=1, cols=1)
                box_table.autofit = False
                cell = box_table.cell(0, 0)
                set_cell_shading(cell, "F2F4F7") # Light Gray Shading
                set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
                
                # Single-cell border
                tcPr = cell._tc.get_or_add_tcPr()
                borders_xml = f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:top w:val="single" w:sz="6" w:space="0" w:color="A6ACAF"/>
                    <w:left w:val="single" w:sz="24" w:space="0" w:color="1F4E79"/>
                    <w:bottom w:val="single" w:sz="6" w:space="0" w:color="A6ACAF"/>
                    <w:right w:val="single" w:sz="6" w:space="0" w:color="A6ACAF"/>
                </w:tcBorders>
                '''
                tcPr.append(parse_xml(borders_xml))
                
                # Add code text inside cell
                cell_p = cell.paragraphs[0]
                cell_p.paragraph_format.space_after = Pt(0)
                cell_p.paragraph_format.line_spacing = 1.0
                
                # Render line-by-line
                code_text = "\n".join(code_lines)
                run = cell_p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                
                # Spacer paragraph after table
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                code_lines = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_lines.append(line.rstrip('\n'))
            i += 1
            continue
            
        # --- Handle Tables ---
        if stripped.startswith("|"):
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            # Table ended, render it
            in_table = False
            
            # Clean and parse rows
            parsed_rows = []
            for t_line in table_lines:
                cells = [c.strip() for c in t_line.split("|")[1:-1]]
                parsed_rows.append(cells)
                
            # Filter divider rows like |---|---|
            parsed_rows = [row for row in parsed_rows if not all(re.match(r'^[-:]+$', cell) for cell in row)]
            
            if parsed_rows:
                num_cols = len(parsed_rows[0])
                num_rows = len(parsed_rows)
                
                word_table = doc.add_table(rows=num_rows, cols=num_cols)
                word_table.autofit = True
                set_table_borders(word_table)
                
                # Populate cells
                for r_idx, row_data in enumerate(parsed_rows):
                    for c_idx, cell_value in enumerate(row_data):
                        # Ensure cell exists (failsafe)
                        if c_idx >= num_cols:
                            continue
                        cell = word_table.cell(r_idx, c_idx)
                        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                        
                        cell_p = cell.paragraphs[0]
                        cell_p.paragraph_format.space_after = Pt(0)
                        cell_p.paragraph_format.line_spacing = 1.15
                        
                        # Style headers vs alternate rows
                        if r_idx == 0:
                            # Primary color header (Dark Blue #1F4E79)
                            set_cell_shading(cell, "1F4E79")
                            run = cell_p.add_run(cell_value)
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White text
                            run.font.name = 'Calibri'
                        else:
                            # Striped rows
                            if r_idx % 2 == 0:
                                set_cell_shading(cell, "F2F4F7") # Light Grey zebra shading
                            add_formatted_text(cell_p, cell_value)
                            
                # Spacer paragraph after table
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
            table_lines = []
            # Do not increment i, let current line be processed normally
            continue
            
        # --- Handle General Markdown Syntax ---
        # 1. H1 Heading (# Heading)
        if stripped.startswith("# "):
            title_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(title_text)
            run.bold = True
            run.font.size = Pt(20)
            run.font.name = 'Calibri Light'
            run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Primary Blue
            
        # 2. H2 Heading (## Heading)
        elif stripped.startswith("## "):
            h2_text = stripped[3:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(h2_text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Calibri Light'
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50) # Dark Charcoal
            
        # 3. H3 Heading (### Heading)
        elif stripped.startswith("### "):
            h3_text = stripped[4:]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            
            run = p.add_run(h3_text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x56, 0x65, 0x73) # Soft Blue/Grey
            
        # 4. Horizontal Rule (---)
        elif stripped == "---":
            # Add thin visual horizontal divider
            rule_table = doc.add_table(rows=1, cols=1)
            cell = rule_table.cell(0, 0)
            set_cell_shading(cell, "D3D3D3")
            cell.height = Pt(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
        # 5. Bullet lists (- or *)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            add_formatted_text(p, bullet_text)
            
        # 6. Blank / Empty lines
        elif not stripped:
            pass # Skip extra blanks to allow Word's paragraph spacing to handle layout
            
        # 7. Normal paragraphs
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)
            
        i += 1
        
    print(f"Saving compiled Word document to {docx_path}...")
    doc.save(docx_path)
    print("Compilation completed successfully!")

if __name__ == "__main__":
    md_proposal = "ColonyAI_Proposal.md"
    docx_proposal = "Proposal_AI_Open_ColonyAI.docx"
    print("ColonyAI Word Document Compiler v1.0")
    convert_md_to_docx(md_proposal, docx_proposal)
